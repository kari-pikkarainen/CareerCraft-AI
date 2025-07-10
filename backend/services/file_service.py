"""
File Upload and Processing Service

Handles secure file uploads, validation, and text extraction for resumes.
Supports PDF, DOCX, and TXT formats with comprehensive security checks.
"""

import os
import tempfile
import logging
from pathlib import Path
from typing import Optional, Dict, Any, List, Tuple
from datetime import datetime, timezone
import secrets
import hashlib
# Try to import magic, fallback to mimetypes if not available
try:
    import magic
    MAGIC_AVAILABLE = True
except ImportError:
    import mimetypes
    MAGIC_AVAILABLE = False
from dataclasses import dataclass

# File processing imports
import PyPDF2
from docx import Document
import chardet

from config import get_config
from utils.parsers import get_resume_parser

logger = logging.getLogger(__name__)

@dataclass
class FileInfo:
    """File information container"""
    file_id: str
    original_filename: str
    file_size: int
    content_type: str
    file_format: str
    upload_timestamp: datetime
    file_hash: str
    temp_path: Optional[Path] = None

@dataclass
class ExtractedContent:
    """Extracted file content container"""
    text: str
    metadata: Dict[str, Any]
    sections: Dict[str, str]
    word_count: int
    extraction_method: str


class FileValidationError(Exception):
    """Raised when file validation fails"""
    pass

class FileProcessingError(Exception):
    """Raised when file processing fails"""
    pass


class FileService:
    """Service for handling file uploads and processing"""
    
    # Supported file types
    SUPPORTED_FORMATS = {
        'pdf': ['application/pdf'],
        'docx': [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/vnd.ms-word.document.macroEnabled.12'
        ],
        'txt': ['text/plain', 'text/x-plain', 'application/octet-stream']
    }
    
    # File size limits (configurable)
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB default
    
    # Security patterns to block
    DANGEROUS_PATTERNS = [
        b'<script',
        b'javascript:',
        b'vbscript:',
        b'onload=',
        b'onerror=',
        b'<?php',
        b'<%',
        b'#!/bin/',
        b'#!/usr/bin/'
    ]
    
    def __init__(self):
        """Initialize file service"""
        self.config = get_config()
        self.temp_dir = Path(tempfile.gettempdir()) / "careercraft_uploads"
        self.temp_dir.mkdir(exist_ok=True)
        
        # Update max file size from config
        if hasattr(self.config.security, 'max_file_size'):
            self.MAX_FILE_SIZE = self.config.security.max_file_size
    
    def validate_file(self, file_content: bytes, filename: str, content_type: str) -> str:
        """
        Comprehensive file validation.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            content_type: MIME content type
            
        Returns:
            Detected file format (pdf, docx, txt)
            
        Raises:
            FileValidationError: If validation fails
        """
        # Check file size
        if len(file_content) == 0:
            raise FileValidationError("Empty file not allowed")
            
        if len(file_content) > self.MAX_FILE_SIZE:
            raise FileValidationError(f"File size exceeds limit of {self.MAX_FILE_SIZE // (1024*1024)}MB")
        
        # Check filename
        if not filename or len(filename) > 255:
            raise FileValidationError("Invalid filename")
            
        # Extract file extension
        file_ext = Path(filename).suffix.lower()
        if file_ext not in ['.pdf', '.docx', '.txt']:
            raise FileValidationError(f"Unsupported file extension: {file_ext}")
        
        # Detect actual file type using magic numbers or fallback to mimetypes
        if MAGIC_AVAILABLE:
            try:
                detected_type = magic.from_buffer(file_content, mime=True)
            except Exception:
                # Fallback: use content type
                detected_type = content_type
        else:
            # Use mimetypes as fallback
            guessed_type, _ = mimetypes.guess_type(filename)
            detected_type = guessed_type or content_type
        
        # Validate file format
        format_name = None
        for fmt, mime_types in self.SUPPORTED_FORMATS.items():
            if detected_type in mime_types:
                format_name = fmt
                break
        
        if not format_name:
            # Special handling for text files
            if file_ext == '.txt' and detected_type.startswith('text/'):
                format_name = 'txt'
            else:
                raise FileValidationError(f"Unsupported file type: {detected_type}")
        
        # Security scan for dangerous content
        self._security_scan(file_content, format_name)
        
        # Format-specific validation
        if format_name == 'pdf':
            self._validate_pdf(file_content)
        elif format_name == 'docx':
            self._validate_docx(file_content)
        elif format_name == 'txt':
            self._validate_text(file_content)
        
        return format_name
    
    def _security_scan(self, content: bytes, file_format: str) -> None:
        """Scan file content for security threats"""
        # Check for dangerous patterns
        content_lower = content.lower()
        for pattern in self.DANGEROUS_PATTERNS:
            if pattern in content_lower:
                raise FileValidationError(f"Suspicious content detected: {pattern.decode('utf-8', errors='ignore')}")
        
        # Additional format-specific security checks
        if file_format == 'pdf':
            # Check for JavaScript in PDF
            if b'/js' in content_lower or b'/javascript' in content_lower:
                raise FileValidationError("PDF contains JavaScript which is not allowed")
        
        elif file_format == 'docx':
            # Check for macros
            if b'vbaProject' in content or b'macros' in content_lower:
                raise FileValidationError("Documents with macros are not allowed")
    
    def _validate_pdf(self, content: bytes) -> None:
        """Validate PDF file structure"""
        try:
            # Create temporary file to validate PDF
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_file:
                temp_file.write(content)
                temp_path = temp_file.name
            
            try:
                # Try to read PDF
                with open(temp_path, 'rb') as pdf_file:
                    reader = PyPDF2.PdfReader(pdf_file)
                    
                    # Check if PDF is readable
                    if len(reader.pages) == 0:
                        raise FileValidationError("PDF contains no readable pages")
                    
                    # Check for excessive page count (potential DoS)
                    if len(reader.pages) > 50:
                        raise FileValidationError("PDF exceeds maximum page limit (50 pages)")
                    
                    # Try to read first page to ensure it's not corrupted
                    try:
                        first_page = reader.pages[0]
                        text = first_page.extract_text()
                        # PDF should have some extractable content
                        if not text or len(text.strip()) < 10:
                            logger.warning("PDF appears to have minimal text content")
                    except Exception:
                        raise FileValidationError("PDF appears to be corrupted or unreadable")
                        
            finally:
                # Cleanup temp file
                os.unlink(temp_path)
                
        except FileValidationError:
            raise
        except Exception as e:
            raise FileValidationError(f"PDF validation failed: {str(e)}")
    
    def _validate_docx(self, content: bytes) -> None:
        """Validate DOCX file structure"""
        try:
            # Create temporary file to validate DOCX
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_file:
                temp_file.write(content)
                temp_path = temp_file.name
            
            try:
                # Try to read DOCX
                doc = Document(temp_path)
                
                # Check if document has content
                if len(doc.paragraphs) == 0:
                    raise FileValidationError("Document contains no readable content")
                
                # Check for excessive content (potential DoS)
                total_text = ''.join([para.text for para in doc.paragraphs])
                if len(total_text) > 1000000:  # 1MB of text
                    raise FileValidationError("Document content exceeds size limit")
                    
            finally:
                # Cleanup temp file
                os.unlink(temp_path)
                
        except FileValidationError:
            raise
        except Exception as e:
            raise FileValidationError(f"DOCX validation failed: {str(e)}")
    
    def _validate_text(self, content: bytes) -> None:
        """Validate text file content"""
        try:
            # Detect encoding
            encoding_result = chardet.detect(content)
            if encoding_result['confidence'] < 0.7:
                raise FileValidationError("Text file encoding could not be reliably detected")
            
            # Try to decode
            text = content.decode(encoding_result['encoding'])
            
            # Check for reasonable content
            if len(text.strip()) < 10:
                raise FileValidationError("Text file appears to be empty or too short")
            
            # Check for binary content in text file
            if '\x00' in text:
                raise FileValidationError("Text file contains binary data")
                
        except UnicodeDecodeError:
            raise FileValidationError("Text file contains invalid characters")
        except FileValidationError:
            raise
        except Exception as e:
            raise FileValidationError(f"Text validation failed: {str(e)}")
    
    def save_file(self, file_content: bytes, filename: str, content_type: str) -> FileInfo:
        """
        Save uploaded file securely.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename
            content_type: MIME content type
            
        Returns:
            FileInfo object with file metadata
        """
        # Validate file
        file_format = self.validate_file(file_content, filename, content_type)
        
        # Generate secure file ID and path
        file_id = f"file_{secrets.token_urlsafe(16)}"
        safe_filename = self._sanitize_filename(filename)
        temp_path = self.temp_dir / f"{file_id}_{safe_filename}"
        
        # Calculate file hash for integrity
        file_hash = hashlib.sha256(file_content).hexdigest()
        
        try:
            # Save file securely
            with open(temp_path, 'wb') as f:
                f.write(file_content)
            
            # Set restrictive permissions
            temp_path.chmod(0o600)
            
            # Create file info
            file_info = FileInfo(
                file_id=file_id,
                original_filename=filename,
                file_size=len(file_content),
                content_type=content_type,
                file_format=file_format,
                upload_timestamp=datetime.now(timezone.utc),
                file_hash=file_hash,
                temp_path=temp_path
            )
            
            logger.info(f"File saved: {file_id} ({file_format}, {len(file_content)} bytes)")
            return file_info
            
        except Exception as e:
            # Cleanup on error
            if temp_path.exists():
                temp_path.unlink()
            raise FileProcessingError(f"Failed to save file: {str(e)}")
    
    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename for secure storage"""
        # Remove path components
        safe_name = Path(filename).name
        
        # Replace dangerous characters
        dangerous_chars = '<>:"/\\|?*'
        for char in dangerous_chars:
            safe_name = safe_name.replace(char, '_')
        
        # Limit length
        if len(safe_name) > 100:
            name_part = Path(safe_name).stem[:90]
            ext_part = Path(safe_name).suffix
            safe_name = f"{name_part}{ext_part}"
        
        return safe_name
    
    def extract_text(self, file_info: FileInfo) -> ExtractedContent:
        """
        Extract text content from uploaded file.
        
        Args:
            file_info: File information object
            
        Returns:
            ExtractedContent with parsed text and metadata
        """
        if not file_info.temp_path or not file_info.temp_path.exists():
            raise FileProcessingError("File not found or already deleted")
        
        try:
            if file_info.file_format == 'pdf':
                return self._extract_pdf_text(file_info)
            elif file_info.file_format == 'docx':
                return self._extract_docx_text(file_info)
            elif file_info.file_format == 'txt':
                return self._extract_text_content(file_info)
            else:
                raise FileProcessingError(f"Unsupported format: {file_info.file_format}")
                
        except Exception as e:
            logger.error(f"Text extraction failed for {file_info.file_id}: {e}")
            raise FileProcessingError(f"Failed to extract text: {str(e)}")
    
    def parse_resume(self, file_info: FileInfo) -> Dict[str, Any]:
        """
        Extract and parse resume with structured data extraction.
        
        Args:
            file_info: File information object
            
        Returns:
            Dictionary with parsed resume data
        """
        # First extract raw text
        extracted_content = self.extract_text(file_info)
        
        # Use advanced parser for structured extraction
        parser = get_resume_parser()
        parsed_resume = parser.parse_resume(
            text=extracted_content.text,
            metadata=extracted_content.metadata
        )
        
        # Convert to dictionary for JSON serialization
        return {
            'file_info': {
                'file_id': file_info.file_id,
                'filename': file_info.original_filename,
                'file_format': file_info.file_format,
                'word_count': extracted_content.word_count
            },
            'contact_info': {
                'email': parsed_resume.contact_info.email,
                'phone': parsed_resume.contact_info.phone,
                'linkedin': parsed_resume.contact_info.linkedin,
                'github': parsed_resume.contact_info.github,
                'website': parsed_resume.contact_info.website,
                'address': parsed_resume.contact_info.address
            },
            'summary': parsed_resume.summary,
            'work_experience': [
                {
                    'company': exp.company,
                    'position': exp.position,
                    'start_date': exp.start_date,
                    'end_date': exp.end_date,
                    'location': exp.location,
                    'description': exp.description,
                    'technologies': exp.technologies
                }
                for exp in parsed_resume.work_experience
            ],
            'education': [
                {
                    'institution': edu.institution,
                    'degree': edu.degree,
                    'field_of_study': edu.field_of_study,
                    'graduation_date': edu.graduation_date,
                    'gpa': edu.gpa,
                    'location': edu.location
                }
                for edu in parsed_resume.education
            ],
            'skills': parsed_resume.skills,
            'projects': [
                {
                    'name': proj.name,
                    'description': proj.description,
                    'technologies': proj.technologies,
                    'url': proj.url
                }
                for proj in parsed_resume.projects
            ],
            'certifications': parsed_resume.certifications,
            'languages': parsed_resume.languages,
            'sections': parsed_resume.sections,
            'extraction_metadata': {
                'extraction_method': extracted_content.extraction_method,
                'word_count': extracted_content.word_count,
                'sections_detected': len(parsed_resume.sections),
                'contact_fields_found': sum(1 for field in [
                    parsed_resume.contact_info.email,
                    parsed_resume.contact_info.phone,
                    parsed_resume.contact_info.linkedin,
                    parsed_resume.contact_info.github
                ] if field),
                'work_experiences_found': len(parsed_resume.work_experience),
                'education_entries_found': len(parsed_resume.education),
                'skills_found': len(parsed_resume.skills),
                'projects_found': len(parsed_resume.projects)
            }
        }
    
    def _extract_pdf_text(self, file_info: FileInfo) -> ExtractedContent:
        """Extract text from PDF file"""
        with open(file_info.temp_path, 'rb') as pdf_file:
            reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            metadata = {
                'page_count': len(reader.pages),
                'pdf_info': {}
            }
            
            # Extract metadata if available
            if reader.metadata:
                metadata['pdf_info'] = {
                    'title': reader.metadata.get('/Title', ''),
                    'author': reader.metadata.get('/Author', ''),
                    'creator': reader.metadata.get('/Creator', ''),
                    'producer': reader.metadata.get('/Producer', ''),
                    'creation_date': str(reader.metadata.get('/CreationDate', ''))
                }
            
            # Extract text from all pages
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num}: {e}")
            
            full_text = '\n'.join(text_parts)
            
            # Basic section detection (simple heuristics)
            sections = self._detect_resume_sections(full_text)
            
            return ExtractedContent(
                text=full_text,
                metadata=metadata,
                sections=sections,
                word_count=len(full_text.split()),
                extraction_method='PyPDF2'
            )
    
    def _extract_docx_text(self, file_info: FileInfo) -> ExtractedContent:
        """Extract text from DOCX file"""
        doc = Document(file_info.temp_path)
        
        text_parts = []
        metadata = {
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
            'core_properties': {}
        }
        
        # Extract core properties if available
        if doc.core_properties:
            metadata['core_properties'] = {
                'author': doc.core_properties.author or '',
                'title': doc.core_properties.title or '',
                'subject': doc.core_properties.subject or '',
                'created': str(doc.core_properties.created) if doc.core_properties.created else '',
                'modified': str(doc.core_properties.modified) if doc.core_properties.modified else ''
            }
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        full_text = '\n'.join(text_parts)
        
        # Basic section detection
        sections = self._detect_resume_sections(full_text)
        
        return ExtractedContent(
            text=full_text,
            metadata=metadata,
            sections=sections,
            word_count=len(full_text.split()),
            extraction_method='python-docx'
        )
    
    def _extract_text_content(self, file_info: FileInfo) -> ExtractedContent:
        """Extract content from text file"""
        with open(file_info.temp_path, 'rb') as f:
            content_bytes = f.read()
        
        # Detect encoding
        encoding_result = chardet.detect(content_bytes)
        encoding = encoding_result['encoding'] or 'utf-8'
        
        try:
            text = content_bytes.decode(encoding)
        except UnicodeDecodeError:
            # Fallback to UTF-8 with error handling
            text = content_bytes.decode('utf-8', errors='replace')
        
        metadata = {
            'encoding': encoding,
            'encoding_confidence': encoding_result['confidence'],
            'line_count': text.count('\n') + 1
        }
        
        # Basic section detection
        sections = self._detect_resume_sections(text)
        
        return ExtractedContent(
            text=text,
            metadata=metadata,
            sections=sections,
            word_count=len(text.split()),
            extraction_method='text'
        )
    
    def _detect_resume_sections(self, text: str) -> Dict[str, str]:
        """
        Basic resume section detection using keywords.
        This is a simple implementation - could be enhanced with ML.
        """
        sections = {}
        text_lines = text.split('\n')
        current_section = 'header'
        current_content = []
        
        # Common section headers
        section_keywords = {
            'experience': ['experience', 'work history', 'employment', 'professional experience', 'career'],
            'education': ['education', 'academic', 'qualification', 'degree', 'university', 'college'],
            'skills': ['skills', 'competencies', 'technical skills', 'expertise', 'proficiencies'],
            'summary': ['summary', 'profile', 'objective', 'about', 'overview'],
            'contact': ['contact', 'phone', 'email', 'address', 'linkedin'],
            'projects': ['projects', 'portfolio', 'achievements', 'accomplishments'],
            'certifications': ['certifications', 'licenses', 'certificates']
        }
        
        for line in text_lines:
            line_lower = line.lower().strip()
            
            # Check if line is a section header
            detected_section = None
            for section_name, keywords in section_keywords.items():
                if any(keyword in line_lower for keyword in keywords):
                    # Additional check: line should be relatively short and not contain too much other text
                    if len(line.strip()) < 100 and any(line_lower.startswith(keyword) or line_lower.endswith(keyword) for keyword in keywords):
                        detected_section = section_name
                        break
            
            if detected_section:
                # Save previous section
                if current_content:
                    sections[current_section] = '\n'.join(current_content).strip()
                
                # Start new section
                current_section = detected_section
                current_content = []
            else:
                # Add to current section
                if line.strip():
                    current_content.append(line)
        
        # Save last section
        if current_content:
            sections[current_section] = '\n'.join(current_content).strip()
        
        return sections
    
    def cleanup_file(self, file_info: FileInfo) -> bool:
        """
        Clean up temporary file.
        
        Args:
            file_info: File information object
            
        Returns:
            True if cleanup successful
        """
        try:
            if file_info.temp_path and file_info.temp_path.exists():
                file_info.temp_path.unlink()
                logger.info(f"Cleaned up file: {file_info.file_id}")
                return True
            return False
        except Exception as e:
            logger.error(f"Failed to cleanup file {file_info.file_id}: {e}")
            return False
    
    def cleanup_old_files(self, max_age_hours: int = 24) -> int:
        """
        Clean up old temporary files.
        
        Args:
            max_age_hours: Maximum age in hours before deletion
            
        Returns:
            Number of files cleaned up
        """
        cleaned_count = 0
        cutoff_time = datetime.now() - datetime.timedelta(hours=max_age_hours)
        
        try:
            for file_path in self.temp_dir.glob('file_*'):
                if file_path.is_file():
                    file_mtime = datetime.fromtimestamp(file_path.stat().st_mtime)
                    if file_mtime < cutoff_time:
                        try:
                            file_path.unlink()
                            cleaned_count += 1
                        except Exception as e:
                            logger.error(f"Failed to delete old file {file_path}: {e}")
            
            if cleaned_count > 0:
                logger.info(f"Cleaned up {cleaned_count} old temporary files")
                
        except Exception as e:
            logger.error(f"Error during cleanup: {e}")
        
        return cleaned_count


# Global file service instance
_file_service_instance = None

def get_file_service() -> FileService:
    """Get file service singleton instance"""
    global _file_service_instance
    if _file_service_instance is None:
        _file_service_instance = FileService()
    return _file_service_instance